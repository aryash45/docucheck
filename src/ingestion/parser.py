"""
src/ingestion/parser.py
-----------------------
DocumentParser — converts PDF, Word (.docx), or plain-text files into a
structured ParsedDocument.

Supported formats
-----------------
  .pdf          PyMuPDF (fitz)   — zero network, zero Java, fast
  .docx         python-docx      — paragraphs + tables
  .txt / .md    built-in open()
  other         chardet + open() binary fallback

Design notes
------------
- Never raises on read failure; returns a ParsedDocument with empty text
  and appends a warning to .warnings so callers can decide what to do.
- All text is normalised: multiple blank lines collapsed, leading/trailing
  whitespace stripped per paragraph.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Try importing optional heavy deps; degrade gracefully ──────────────────

try:
    import fitz  # PyMuPDF
    _HAVE_FITZ = True
except ImportError:
    _HAVE_FITZ = False
    logger.warning("PyMuPDF (fitz) not installed — PDF parsing disabled")

try:
    import docx  # python-docx
    _HAVE_DOCX = True
except ImportError:
    _HAVE_DOCX = False
    logger.warning("python-docx not installed — .docx parsing disabled")

try:
    import chardet
    _HAVE_CHARDET = True
except ImportError:
    _HAVE_CHARDET = False


# ---------------------------------------------------------------------------
# Return type
# ---------------------------------------------------------------------------

@dataclass
class ParsedDocument:
    """Structured representation of an ingested document."""

    text: str                         # Full normalised plain text
    title: str = ""                   # Best-guess title (first non-empty line or filename)
    source_path: str = ""             # Original file path
    page_count: int = 0               # 0 if unknown (e.g. plain text)
    pages: list[str] = field(default_factory=list)  # Per-page text (may be empty)
    warnings: list[str] = field(default_factory=list)

    @property
    def word_count(self) -> int:
        return len(self.text.split())

    @property
    def is_empty(self) -> bool:
        return len(self.text.strip()) < 10


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

class DocumentParser:
    """
    Parse a file from disk into a ParsedDocument.

    Usage
    -----
        parser = DocumentParser()
        doc = parser.parse("paper.pdf")
        print(doc.title, doc.word_count)
    """

    def parse(self, path: str | Path) -> ParsedDocument:
        path = Path(path)
        ext = path.suffix.lower()

        if not path.exists():
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=[f"File not found: {path}"],
            )

        logger.info("DocumentParser: parsing %s (%s)", path.name, ext)

        if ext == ".pdf":
            return self._parse_pdf(path)
        elif ext == ".docx":
            return self._parse_docx(path)
        elif ext in (".txt", ".md", ".rst", ".tex"):
            return self._parse_text(path)
        else:
            return self._parse_binary_fallback(path)

    # ------------------------------------------------------------------
    # PDF
    # ------------------------------------------------------------------

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        if not _HAVE_FITZ:
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=["pymupdf not installed; cannot parse PDF. Run: pip install pymupdf"],
            )
        try:
            doc = fitz.open(str(path))
            pages: list[str] = []
            for page in doc:
                text = page.get_text("text")  # type: ignore[attr-defined]
                pages.append(_normalise(text))
            doc.close()

            full_text = "\n\n".join(p for p in pages if p)
            title = _extract_title(pages[0] if pages else "", path.stem)

            return ParsedDocument(
                text=full_text,
                title=title,
                source_path=str(path),
                page_count=len(pages),
                pages=pages,
            )
        except Exception as exc:
            logger.error("DocumentParser: PDF parse failed: %s", exc)
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=[f"PDF parse error: {exc}"],
            )

    # ------------------------------------------------------------------
    # Word (.docx)
    # ------------------------------------------------------------------

    def _parse_docx(self, path: Path) -> ParsedDocument:
        if not _HAVE_DOCX:
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=["python-docx not installed; cannot parse .docx. Run: pip install python-docx"],
            )
        try:
            document = docx.Document(str(path))
            paragraphs: list[str] = []

            # Body paragraphs
            for para in document.paragraphs:
                stripped = para.text.strip()
                if stripped:
                    paragraphs.append(stripped)

            # Tables
            for table in document.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)
            title = _extract_title(paragraphs[0] if paragraphs else "", path.stem)

            return ParsedDocument(
                text=full_text,
                title=title,
                source_path=str(path),
                page_count=0,   # python-docx doesn't expose page count reliably
                pages=[],
            )
        except Exception as exc:
            logger.error("DocumentParser: DOCX parse failed: %s", exc)
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=[f"DOCX parse error: {exc}"],
            )

    # ------------------------------------------------------------------
    # Plain text / Markdown / LaTeX / RST
    # ------------------------------------------------------------------

    def _parse_text(self, path: Path) -> ParsedDocument:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            normalised = _normalise(text)
            first_line = next((l for l in normalised.splitlines() if l.strip()), path.stem)
            title = first_line.strip().lstrip("#").strip()[:120]
            return ParsedDocument(
                text=normalised,
                title=title,
                source_path=str(path),
                page_count=0,
            )
        except Exception as exc:
            logger.error("DocumentParser: text parse failed: %s", exc)
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=[f"Text parse error: {exc}"],
            )

    # ------------------------------------------------------------------
    # Unknown binary fallback
    # ------------------------------------------------------------------

    def _parse_binary_fallback(self, path: Path) -> ParsedDocument:
        """Attempt chardet encoding detection and read as text."""
        raw = path.read_bytes()
        encoding = "utf-8"
        if _HAVE_CHARDET:
            detected = chardet.detect(raw)
            encoding = detected.get("encoding") or "utf-8"
        try:
            text = raw.decode(encoding, errors="replace")
            normalised = _normalise(text)
            return ParsedDocument(
                text=normalised,
                title=path.stem,
                source_path=str(path),
                warnings=[f"Unknown extension '{path.suffix}'; decoded as {encoding}"],
            )
        except Exception as exc:
            return ParsedDocument(
                text="",
                source_path=str(path),
                warnings=[f"Binary fallback failed: {exc}"],
            )


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalise(text: str) -> str:
    """Collapse excessive blank lines and strip per-line trailing whitespace."""
    lines = text.splitlines()
    out: list[str] = []
    blank_count = 0
    for line in lines:
        stripped = line.rstrip()
        if not stripped:
            blank_count += 1
            if blank_count <= 2:
                out.append("")
        else:
            blank_count = 0
            out.append(stripped)
    return "\n".join(out).strip()


def _extract_title(first_page_text: str, filename_stem: str) -> str:
    """
    Heuristic: the document title is the first non-trivially-short line
    of the first page (or the filename stem as fallback).
    """
    for line in first_page_text.splitlines():
        stripped = line.strip().lstrip("#").strip()
        if len(stripped) > 8:
            return stripped[:150]
    return filename_stem
